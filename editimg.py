from docx import Document  #biblioteca que importa documentos  pip install python-docx
from docx.shared import Inches

def editimg():
    # Carregue o arquivo Word existente
    document = Document("MEMORIAL.docx")

    # Itere pelas tabelas do doc
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if 'Fig1' in run.text:
                            # Substitua pelo caminho da nova imagem
                            for existing_run in paragraph.runs:
                                paragraph._element.remove(existing_run._element)
                            # Adicione a nova imagem (substitua pelo caminho correto)
                            paragraph.add_run().add_picture('fig1.png',width=Inches(5))
                        if 'Fig2' in run.text:
                            # Substitua pelo caminho da nova imagem
                            for existing_run in paragraph.runs:
                                paragraph._element.remove(existing_run._element)
                            # Adicione a nova imagem (substitua pelo caminho correto)
                            paragraph.add_run().add_picture('fig2.png',width=Inches(4))
                        if 'Fig3' in run.text:
                            # Substitua pelo caminho da nova imagem
                            for existing_run in paragraph.runs:
                                paragraph._element.remove(existing_run._element)
                            # Adicione a nova imagem (substitua pelo caminho correto)
                            paragraph.add_run().add_picture('fig3.png',width=Inches(4))
                        if 'Fig4' in run.text:
                            # Substitua pelo caminho da nova imagem
                            for existing_run in paragraph.runs:
                                paragraph._element.remove(existing_run._element)
                            # Adicione a nova imagem (substitua pelo caminho correto)
                            paragraph.add_run().add_picture('fig4.png',width=Inches(3))    

    # Salve o documento modificado
    document.save("CACETA.docx")
    print("imagens subistituida")